#!/usr/bin/env python3
"""
Tilbud Generator
Extracts project info from PDF, Excel, JPEG/PNG files using Claude Vision,
then generates a ready-to-use Ferro tilbud as .docx and pre-filled HTML.

Usage:
  python3 tilbud_generator.py file1.pdf file2.xlsx drawing.png
  python3 tilbud_generator.py --folder ./prosjekt_mappe
"""

import sys
import os
import base64
import json
import re
from pathlib import Path
from datetime import date

# ── third-party ──────────────────────────────────────────────────────────────
import anthropic
import fitz                          # PyMuPDF
import openpyxl
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import io

# ─────────────────────────────────────────────────────────────────────────────
# CONFIG
# ─────────────────────────────────────────────────────────────────────────────
MODEL = "claude-sonnet-4-6"
MAX_IMAGE_DIM = 1568          # Claude vision limit (longest side)
PDF_DPI = 150                 # render resolution for PDF pages

SCOPE_KEYS = [
    "bygget", "grunn", "betong", "staal", "yttervegg", "innervegg",
    "himling", "gulv", "trapper", "vinduer", "porter", "tak",
    "inventar", "ventilasjon", "brann", "prosjektering", "mengder", "rigg",
]

SCOPE_LABELS = {
    "bygget":        "Bygget",
    "grunn":         "Grunnarbeid",
    "betong":        "Betongarbeid",
    "staal":         "Stål",
    "yttervegg":     "Yttervegger",
    "innervegg":     "Innervegg",
    "himling":       "Himling innvendig",
    "gulv":          "Overflater gulver",
    "trapper":       "Trapper / Lås og beslag",
    "vinduer":       "Vinduer og dører",
    "porter":        "Porter",
    "tak":           "Tak",
    "inventar":      "Inventar / Foliering / Screens / Rørlegger / Elektro",
    "ventilasjon":   "Ventilasjon",
    "brann":         "Branntetting/sikring",
    "prosjektering": "Prosjektering / Arkitekt",
    "mengder":       "Mengder",
    "rigg":          "Rigg og drift",
}

FORUTSETNINGER = """Tiltakshaver står selv for:
- Grunnarbeid og komprimering av byggetomt
- Offentlige gebyrer og søknadskostnader
- Fiber inn til bygning
- Dokumentasjon av grunnforhold og uavhengig kontroll
- Brannisolering - tilbys separat ved behov
- Klimagass- og bygningsfysikkrapport - tilbys separat ved behov
- All rivning og klargjøring av byggeplass
- Oljeutskiller og tilknyttet rørarbeid
- Trær, busker og inventar

Generelle betingelser:
- Ved tilleggsarbeid: 750,- pr. time for montør, 1 200,- pr. time for prosjektleder, 15 % materialpåslag.
- Mengder gjeldende og reguleres før kontrakt i samråd med tiltakshaver.
- Tilbudet skal skriftlig bestilles av kunde.
- Medfølgende tegning på stålkonstruksjon gjelder for pris. Det forutsettes kontinuerlig montasje.
- Prisene er basert på anbudsdagens materialpriser og lønninger med vanlig adgang til forhåndsmessig prisjustering.
- Fundamentering dimensjonert for monteringslaster. Setninger fra ferdigstilt grunnarbeid er ikke Ferro sitt ansvar.
- Tiltakshaver sørger for fremkommelig vei rundt bygget - minimum 4 m bredde for kran og transport.
- Tilbudet er gyldig i 14 dager.
- Tilbudet på stål er bygd på gårsdagens innkjøpspriser av stålprofiler. Verkene har varslet prisoppgang og holder kun prisene på dagsbasis. Vi forbeholder oss retten til gjennomgang og eventuell endring av innkjøpspriser ved kontrahering. Endringen gjelder kun stålprofil innkjøp.
- Forutsetter tiltaksklasse 2. Seismikk utelates. Direkte fundamentering 250 kN/m² bruddgrense (fjell og sprengstein). Bygget forutsettes ikke pelet.
- U-verdi tak 0,18 - U-verdi vegg 0,18 - U-verdi glass 1,2.
- Tilbudet er gitt under en pågående krig i Ukraina - ansett som force majeure iht. NS 8417 pkt. 33 / NS 8415 pkt. 24. Fremdriftshindring som forårsakes av krigen gir krav på fristforlengelse og eventuelt tilleggsvederlag.
- Ryddet ut etter eget arbeid - ikke vasket. Offentlig vann/avløp og strømtilførsel til byggegrop er med i pris."""


# ─────────────────────────────────────────────────────────────────────────────
# FILE READERS
# ─────────────────────────────────────────────────────────────────────────────

def _resize_image_bytes(img_bytes: bytes) -> bytes:
    """Scale image so longest side ≤ MAX_IMAGE_DIM, return JPEG bytes."""
    img = Image.open(io.BytesIO(img_bytes)).convert("RGB")
    w, h = img.size
    longest = max(w, h)
    if longest > MAX_IMAGE_DIM:
        scale = MAX_IMAGE_DIM / longest
        img = img.resize((int(w * scale), int(h * scale)), Image.LANCZOS)
    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=85)
    return buf.getvalue()


def read_image_file(path: Path) -> list[dict]:
    """Return list of Claude image content blocks from an image file."""
    raw = path.read_bytes()
    data = base64.standard_b64encode(_resize_image_bytes(raw)).decode()
    return [{"type": "image", "source": {"type": "base64", "media_type": "image/jpeg", "data": data}}]


def read_pdf_file(path: Path) -> list[dict]:
    """Return list of Claude content blocks: text + images for each PDF page."""
    blocks = []
    doc = fitz.open(str(path))
    for page_num, page in enumerate(doc):
        # Extract text first (fast, cheap)
        text = page.get_text().strip()
        if text:
            blocks.append({"type": "text", "text": f"[PDF side {page_num+1} tekst]\n{text}"})
        # Render page as image (catches drawings, tables as images)
        mat = fitz.Matrix(PDF_DPI / 72, PDF_DPI / 72)
        pix = page.get_pixmap(matrix=mat, colorspace=fitz.csRGB)
        img_bytes = pix.tobytes("jpeg")
        img_bytes = _resize_image_bytes(img_bytes)
        data = base64.standard_b64encode(img_bytes).decode()
        blocks.append({
            "type": "image",
            "source": {"type": "base64", "media_type": "image/jpeg", "data": data}
        })
    doc.close()
    return blocks


def read_excel_file(path: Path) -> list[dict]:
    """Extract all cell values from Excel and return as text block."""
    wb = openpyxl.load_workbook(str(path), data_only=True)
    lines = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        lines.append(f"[Excel ark: {sheet_name}]")
        for row in ws.iter_rows(values_only=True):
            row_vals = [str(c) if c is not None else "" for c in row]
            line = "\t".join(row_vals).strip()
            if line.replace("\t", ""):
                lines.append(line)
    return [{"type": "text", "text": "\n".join(lines)}]


def read_docx_file(path: Path) -> list[dict]:
    """Extract all text from a .docx file."""
    doc = Document(str(path))
    lines = []
    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                lines.append(row_text)
    return [{"type": "text", "text": f"[Word-dokument: {path.name}]\n" + "\n".join(lines)}]


def collect_content_blocks(files: list[Path]) -> list[dict]:
    """Read all input files and return merged list of Claude content blocks."""
    blocks = []
    for f in files:
        ext = f.suffix.lower()
        print(f"  Leser: {f.name}")
        if ext == ".pdf":
            blocks.extend(read_pdf_file(f))
        elif ext in (".jpg", ".jpeg", ".png"):
            blocks += [{"type": "text", "text": f"[Bilde: {f.name}]"}]
            blocks.extend(read_image_file(f))
        elif ext in (".xlsx", ".xls"):
            blocks.extend(read_excel_file(f))
        elif ext in (".docx", ".doc"):
            blocks.extend(read_docx_file(f))
        else:
            print(f"    (ukjent filtype, hopper over)")
    return blocks


# ─────────────────────────────────────────────────────────────────────────────
# CLAUDE EXTRACTION
# ─────────────────────────────────────────────────────────────────────────────

EXTRACTION_SCHEMA = {
    "prosjekt":   "Prosjektnavn / byggnavn (string)",
    "kunde":      "Firma/kunde navn (string)",
    "kontakt":    "Kontaktperson (string)",
    "adresse":    "Gateadresse (string)",
    "postnr":     "Postnummer (string)",
    "sted":       "Poststed (string)",
    "sum_eks_mva": "Total pris eks. mva i NOK som tall (number or null)",
    "opsjoner_eks_mva": "Sum opsjoner eks. mva i NOK som tall (number or null)",
    "opsjoner_liste": "Liste med opsjon-tekster og priser (list of strings)",
    "scope": {
        k: f"Beskrivelse for '{SCOPE_LABELS[k]}' (string, tom om ikke nevnt)"
        for k in SCOPE_KEYS
    }
}

SYSTEM_PROMPT = """Du er ekspert på å lese tekniske tegninger, kalkyler og anbudsdokumenter for norske stålentreprenører.
Returner KUN gyldig JSON, ingen forklaring, ingen markdown-blokker."""

USER_PROMPT = f"""Analyser alle vedlagte filer (tegninger, kalkyler, Excel-ark, bilder).
Trekk ut all informasjon du finner og fyll ut denne JSON-strukturen så fullstendig som mulig:

{json.dumps(EXTRACTION_SCHEMA, ensure_ascii=False, indent=2)}

Regler:
- sum_eks_mva: kun tall uten mellomrom/komma/valuta, f.eks. 1234567
- scope-felter: skriv fullstendig beskrivende tekst basert på tegning/kalkyle, tom streng om ikke funnet
- bygget-feltet: beskriv bygget (type, mål, høyde, tak, fasade) fra tegning
- Svar KUN med JSON-objektet."""


def extract_with_claude(content_blocks: list[dict]) -> dict:
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        print("\nFEIL: ANTHROPIC_API_KEY er ikke satt.")
        print("Sett den med:  export ANTHROPIC_API_KEY='sk-ant-...'")
        sys.exit(1)

    client = anthropic.Anthropic(api_key=api_key)

    print("\nSender til Claude for analyse...")
    response = client.messages.create(
        model=MODEL,
        max_tokens=4096,
        system=SYSTEM_PROMPT,
        messages=[{
            "role": "user",
            "content": content_blocks + [{"type": "text", "text": USER_PROMPT}]
        }]
    )

    raw = response.content[0].text.strip()
    # Strip markdown code fences if present
    raw = re.sub(r"^```json\s*", "", raw)
    raw = re.sub(r"\s*```$", "", raw)

    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        print("ADVARSEL: Claude returnerte ugyldig JSON. Prøver å reparere...")
        # Try to extract JSON object
        match = re.search(r"\{.*\}", raw, re.DOTALL)
        if match:
            return json.loads(match.group())
        print("Kunne ikke parse JSON. Claude-svar:")
        print(raw[:500])
        sys.exit(1)


# ─────────────────────────────────────────────────────────────────────────────
# FORMATERING
# ─────────────────────────────────────────────────────────────────────────────

def fmt_kr(n) -> str:
    if n is None:
        return "000 000,-"
    try:
        v = int(round(float(n) / 1000) * 1000)
        return f"{v:,},-".replace(",", " ")
    except (ValueError, TypeError):
        return "000 000,-"

def fmt_exact(n) -> str:
    if n is None:
        return "000 000,-"
    try:
        v = int(round(float(n)))
        return f"{v:,},-".replace(",", " ")
    except (ValueError, TypeError):
        return "000 000,-"


# ─────────────────────────────────────────────────────────────────────────────
# DOCX GENERATOR
# ─────────────────────────────────────────────────────────────────────────────

def _add_heading(doc, text, level=1):
    p = doc.add_paragraph(text)
    run = p.runs[0]
    run.bold = True
    run.font.size = Pt(11 if level == 1 else 10)
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(2)
    return p


def _add_body(doc, text):
    if not text:
        return
    for line in text.split("\n"):
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(1)
        for run in p.runs:
            run.font.size = Pt(10)


def generate_docx(data: dict, output_path: Path):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    today_str = date.today().strftime("%d.%m.%Y")
    prosjekt = data.get("prosjekt") or "Prosjekt"

    # Title
    title = doc.add_paragraph(f"PRISTILBUD – {prosjekt}")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(14)
    doc.add_paragraph(f"Dato: {today_str}")
    doc.add_paragraph()

    # Recipient
    _add_heading(doc, "TILBUD TIL", 1)
    for field in ["kunde", "kontakt", "adresse"]:
        val = data.get(field, "")
        if field == "kontakt" and val:
            val = f"v/ {val}"
        if val:
            doc.add_paragraph(val)
    postnr = data.get("postnr", "")
    sted = data.get("sted", "")
    if postnr or sted:
        doc.add_paragraph(f"{postnr} {sted}".strip())
    doc.add_paragraph()

    # Price summary
    sum_eks = data.get("sum_eks_mva")
    try:
        sum_eks_f = float(sum_eks) if sum_eks else 0
    except (ValueError, TypeError):
        sum_eks_f = 0
    sum_eks_r = round(sum_eks_f / 1000) * 1000
    mva = round(sum_eks_r * 0.25)
    sum_ink = sum_eks_r + mva

    _add_heading(doc, "PRISSAMMENDRAG", 1)
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    rows_data = [
        ("Sum uten opsjoner eks. mva", fmt_kr(sum_eks_r) if sum_eks_r else "000 000,-"),
        ("Mva (25 %)",                 fmt_exact(mva) if sum_eks_r else "000 000,-"),
        ("Sum ink. mva",               fmt_exact(sum_ink) if sum_eks_r else "000 000,-"),
        (f"Gyldig i 14 dager fra {today_str}", ""),
    ]
    for i, (label, val) in enumerate(rows_data):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = val
        if i == 2:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True

    doc.add_paragraph()
    doc.add_paragraph(
        "Vi takker for deres forespørsel og har gleden av å tilby dere følgende:\n"
        "I vedleggene ligger vår beskrivelse av arbeidet. Vi håper tilbudet er tilfredsstillende "
        "og at deres ønsker har blitt ivaretatt. Håper dere aksepterer tilbudet og tar kontakt ved spørsmål."
    )
    doc.add_paragraph()
    doc.add_paragraph("Med vennlig hilsen")
    doc.add_paragraph("Vegard Landsdal")
    doc.add_paragraph("Daglig leder, Ferro Stålentreprenør AS")
    doc.add_paragraph("vegard@ferrostal.no  •  Tlf: 95 83 71 23")

    # Work description
    doc.add_paragraph()
    doc.add_paragraph("─" * 50)
    _add_heading(doc, "ARBEIDSBESKRIVELSE", 1)
    doc.add_paragraph("─" * 50)

    scope = data.get("scope", {})
    for key in SCOPE_KEYS:
        text = scope.get(key, "")
        if text:
            _add_heading(doc, SCOPE_LABELS[key].upper(), 2)
            _add_body(doc, text)

    # Opsjoner
    ops_val = data.get("opsjoner_eks_mva")
    ops_liste = data.get("opsjoner_liste", [])
    try:
        ops_f = float(ops_val) if ops_val else 0
    except (ValueError, TypeError):
        ops_f = 0

    if ops_f > 0 or ops_liste:
        doc.add_paragraph()
        _add_heading(doc, "OPSJONER (eks. mva)", 1)
        for item in (ops_liste or []):
            doc.add_paragraph(f"• {item}")
        if ops_f > 0:
            ops_mva = round(ops_f * 0.25)
            ops_ink = ops_f + ops_mva
            doc.add_paragraph()
            doc.add_paragraph(f"Sum opsjoner eks. mva:   {fmt_exact(ops_f)}")
            doc.add_paragraph(f"Mva (25 %):              {fmt_exact(ops_mva)}")
            doc.add_paragraph(f"Sum opsjoner ink. mva:   {fmt_exact(ops_ink)}")

    # Forutsetninger
    doc.add_paragraph()
    doc.add_paragraph("─" * 50)
    _add_heading(doc, "GENERELLE FORUTSETNINGER", 1)
    doc.add_paragraph("─" * 50)
    _add_body(doc, FORUTSETNINGER)

    doc.add_paragraph()
    doc.add_paragraph("Ferro Stålentreprenør AS  •  Ringsevja 3, 3830 Ulefoss  •  Tlf: 95 83 71 23")

    doc.save(str(output_path))


# ─────────────────────────────────────────────────────────────────────────────
# HTML GENERATOR (pre-filled tilbud_html.html clone)
# ─────────────────────────────────────────────────────────────────────────────

def _js_str(s) -> str:
    """Escape string for JavaScript template literal."""
    if not s:
        return ""
    return s.replace("\\", "\\\\").replace("`", "\\`").replace("${", "\\${").replace("\n", "\\n")


def generate_prefilled_html(data: dict, output_path: Path):
    scope = data.get("scope", {})
    scope_js = json.dumps({k: scope.get(k, "") for k in SCOPE_KEYS}, ensure_ascii=False)

    info_js = json.dumps({
        "prosjekt": data.get("prosjekt", ""),
        "kunde":    data.get("kunde", ""),
        "kontakt":  data.get("kontakt", ""),
        "adresse":  data.get("adresse", ""),
        "postnr":   data.get("postnr", ""),
        "sted":     data.get("sted", ""),
        "dato":     date.today().isoformat(),
    }, ensure_ascii=False)

    sum_eks = data.get("sum_eks_mva", "")
    ops_eks = data.get("opsjoner_eks_mva", "")

    # Read original HTML and inject initial state
    orig = Path(__file__).parent / "tilbud_html.html"
    html = orig.read_text(encoding="utf-8")

    # Inject pre-fill script just before </body>
    inject = f"""
<script type="text/babel">
// Pre-fill data injected by tilbud_generator.py
window.__PREFILL__ = {{
  info: {info_js},
  scope: {scope_js},
  sumEks: {json.dumps(str(sum_eks) if sum_eks else "")},
  opsjoner: {json.dumps(str(ops_eks) if ops_eks else "")},
}};
</script>
"""
    # Patch useState in original script to pick up prefill
    patch = """
  const [tab, setTab] = useState("info");
  const pf = window.__PREFILL__ || {};
  const [info, setInfo] = useState(pf.info || { prosjekt:"", kunde:"", kontakt:"", adresse:"", postnr:"", sted:"", dato: today() });
  const [pris, setPris] = useState({ sumEks: pf.sumEks||"", opsjoner: pf.opsjoner||"" });
  const [scope, setScope] = useState(() => { const s={}; SCOPE_ITEMS.forEach(i=>{ s[i.key]=(pf.scope&&pf.scope[i.key]!==undefined)?pf.scope[i.key]:i.def; }); return s; });"""

    original_state = """
  const [tab, setTab] = useState("info");
  const [info, setInfo] = useState({ prosjekt:"", kunde:"", kontakt:"", adresse:"", postnr:"", sted:"", dato: today() });
  const [pris, setPris] = useState({ sumEks:"", opsjoner:"" });
  const [scope, setScope] = useState(() => { const s={}; SCOPE_ITEMS.forEach(i=>{ s[i.key]=i.def; }); return s; });"""

    html = html.replace(original_state, patch)
    html = html.replace("</body>", inject + "\n</body>")

    output_path.write_text(html, encoding="utf-8")


# ─────────────────────────────────────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────────────────────────────────────

def collect_files(args: list[str]) -> list[Path]:
    files = []
    i = 0
    while i < len(args):
        arg = args[i]
        if arg == "--folder" and i + 1 < len(args):
            folder = Path(args[i + 1])
            for ext in ("*.pdf", "*.xlsx", "*.xls", "*.jpg", "*.jpeg", "*.png"):
                files.extend(sorted(folder.glob(ext)))
            i += 2
        else:
            p = Path(arg)
            if p.exists():
                files.append(p)
            else:
                print(f"ADVARSEL: Fant ikke filen: {arg}")
            i += 1
    return files


def main():
    args = sys.argv[1:]
    if not args:
        print(__doc__)
        sys.exit(0)

    files = collect_files(args)
    if not files:
        print("Ingen gyldige filer funnet.")
        sys.exit(1)

    print(f"\nFunnet {len(files)} fil(er):")
    for f in files:
        print(f"  {f}")

    print("\nLeser filer...")
    content_blocks = collect_content_blocks(files)

    data = extract_with_claude(content_blocks)

    print("\nEkstrahert data:")
    print(f"  Prosjekt : {data.get('prosjekt', '-')}")
    print(f"  Kunde    : {data.get('kunde', '-')}")
    print(f"  Sum eks. : {data.get('sum_eks_mva', '-')}")
    filled = sum(1 for k in SCOPE_KEYS if data.get("scope", {}).get(k))
    print(f"  Scope    : {filled}/{len(SCOPE_KEYS)} felter fylt")

    # Output in same folder as first input file
    out_dir = files[0].parent
    prosjekt_slug = re.sub(r"[^\w\-]", "_", data.get("prosjekt") or "Tilbud")
    today_str = date.today().isoformat()
    base_name = f"Tilbud_{prosjekt_slug}_{today_str}"

    docx_path = out_dir / f"{base_name}.docx"
    html_path = out_dir / f"{base_name}_prefilled.html"

    print(f"\nGenererer .docx  → {docx_path.name}")
    generate_docx(data, docx_path)

    print(f"Genererer HTML   → {html_path.name}")
    generate_prefilled_html(data, html_path)

    print("\nFerdig!")
    print(f"  {docx_path}")
    print(f"  {html_path}")

    # Also save extracted JSON for debugging
    json_path = out_dir / f"{base_name}_extracted.json"
    json_path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"  {json_path}  (rådata fra Claude)")


if __name__ == "__main__":
    main()
